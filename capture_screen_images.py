import tkinter as tk
from tkinter import Label, StringVar, filedialog, Checkbutton, BooleanVar, simpledialog, Entry, messagebox
from PIL import Image, ImageGrab
from docx import Document
import time
import threading
import pyautogui

class CaptureWindow:
    def __init__(self, root):
        self.root = root
        self.root.title("Capture Window")
        self.root.attributes("-topmost", True)

        # Entry for trimming dimensions
        self.height_label = Label(root, text="Trim Height (mm):")
        self.height_label.pack(pady=5)
        self.height_entry = Entry(root)
        self.height_entry.pack()

        self.width_label = Label(root, text="Trim Width (mm):")
        self.width_label.pack(pady=5)
        self.width_entry = Entry(root)
        self.width_entry.pack()

        # Buttons for capture and browsing
        self.capture_button = tk.Button(root, text="Capture", command=self.capture_action)
        self.capture_button.pack(pady=10)

        self.browse_button = tk.Button(root, text="Browse", command=self.browse_action)
        self.browse_button.pack(pady=10)

        # Path display
        self.path_var = StringVar()
        self.path_var.set("C:/temp/")
        self.path_label = Label(root, text=f"Path: {self.path_var.get()}screenshot.docx")
        self.path_label.pack()

        self.document = Document()

        # Checkboxes for additional options
        self.comment_var = BooleanVar()
        self.comment_var.set(False)
        self.comment_checkbox = Checkbutton(root, text="Ask for input comments", variable=self.comment_var)
        self.comment_checkbox.pack(pady=5)

        self.timer_var = BooleanVar()
        self.timer_var.set(False)
        self.timer_checkbox = Checkbutton(root, text="Enable Timer", variable=self.timer_var, command=self.toggle_timer)
        self.timer_checkbox.pack(pady=5)

        self.timer_label = Label(root, text="Timer: ")
        self.timer_label.pack()

        # Custom capture button
        self.custom_capture_button = tk.Button(root, text="Custom Capture", command=self.custom_capture_action)
        self.custom_capture_button.pack(pady=10)

        # X and Y coordinates entry
        self.x_label = Label(root, text="X Coordinate:")
        self.x_label.pack(pady=5)
        self.x_entry = Entry(root)
        self.x_entry.pack()

        self.y_label = Label(root, text="Y Coordinate:")
        self.y_label.pack(pady=5)
        self.y_entry = Entry(root)
        self.y_entry.pack()

        # Other attributes
        self.capture_phase = 1
        self.x_screen = 0
        self.y_screen = 0

        self.root.title("Capture Window")
        self.root.attributes("-topmost", True)

        # Labels to display mouse coordinates
        self.x_label = Label(root, text="X Coordinate: ")
        self.x_label.pack(pady=5)

        self.y_label = Label(root, text="Y Coordinate: ")
        self.y_label.pack(pady=5)

        # Update coordinates continuously
        self.update_coordinates()

    def toggle_timer(self):
        if self.timer_var.get():
            self.timer_thread = threading.Thread(target=self.start_timer)
            self.timer_thread.start()
        else:
            self.timer_label.config(text="Timer: ")

    def start_timer(self):
        for i in range(10, 0, -1):
            self.timer_label.config(text=f"Timer: {i}")
            time.sleep(1)
        self.capture_action()

    def validate_input(self):
        height_input = self.height_entry.get()
        width_input = self.width_entry.get()

        # If height or width is empty, return True to proceed with capturing
        if not height_input and not width_input:
            return True

        try:
            height_mm = float(height_input) if height_input else None
            width_mm = float(width_input) if width_input else None

            if height_mm is not None and height_mm <= 0 or width_mm is not None and width_mm <= 0:
                messagebox.showerror("Error", "Height and width must be positive numbers.")
                return False
            return True
        except ValueError:
            messagebox.showerror("Error", "Please enter valid numeric values for height and width.")
            return False

    def capture_action(self):
        print("Capturing the entire screen.")

        # Validate input
        if not self.validate_input():
            return

        # Hide the window temporarily
        self.root.withdraw()
        time.sleep(1)  # Wait for 1 second

        # Get the specified height and width
        height_mm = float(self.height_entry.get()) if self.height_entry.get() else None
        width_mm = float(self.width_entry.get()) if self.width_entry.get() else None

        # Convert mm to pixels
        if height_mm:
            height_px = int(height_mm * 3.7795275591)  # 1 mm = 3.7795275591 pixels
        else:
            height_px = 0

        if width_mm:
            width_px = int(width_mm * 3.7795275591)
        else:
            width_px = 0

        # Capture the screenshot of the entire screen
        screenshot = ImageGrab.grab()

        # Trim specified amount from top and bottom
        top_trim = height_px
        bottom_trim = height_px
        left_trim = 0
        right_trim = 0
        screenshot = screenshot.crop((left_trim, top_trim, screenshot.width - right_trim, screenshot.height - bottom_trim))

        # Save the screenshot to a file
        screenshot_path = f"{self.path_var.get()}Screenshot.docx"
        screenshot.save(f"{self.path_var.get()}Screenshot_temp.png")

        # Add the screenshot to the Word document
        self.document.add_picture(f"{self.path_var.get()}screenshot_temp.png")

        # Prompt for input comments if the checkbox is checked
        if self.comment_var.get():
            comment = simpledialog.askstring("Input", "Enter your comment:")
            if comment:
                self.document.add_paragraph(comment)

        # Save the Word document
        self.document.save(screenshot_path)

        print(f"Screenshot captured and added to {screenshot_path}")

        # Update the path label
        self.path_label.config(text=f"Path: {screenshot_path}")

        # Show the window again
        self.root.deiconify()

    def browse_action(self):
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.path_var.set(folder_path + "/")
            print(f"Destination folder set to: {self.path_var.get()}")

    def custom_capture_action(self):
        print("Capturing custom part of the screen.")

        # Validate input
        if not self.validate_input():
            return

        # Set up a timer to change the capture phase after 5 seconds
        threading.Timer(5, self.capture_second_phase).start()

    def capture_second_phase(self):
        # Change the capture phase after 5 seconds
        self.capture_phase = 2

        # Get the coordinates for the first phase
        self.x1_coordinate = self.x_screen
        self.y1_coordinate = self.y_screen

        # Prompt the user to capture the second set of coordinates
        messagebox.showinfo("Capture Second Phase", "Please move your mouse to capture the second set of coordinates.")

        # Schedule capturing after 5 more seconds
        threading.Timer(5, self.capture_custom_screenshot).start()

    def capture_custom_screenshot(self):
        # Get the second set of coordinates
        x2_coordinate = self.x_screen
        y2_coordinate = self.y_screen

        # Hide the window temporarily
        self.root.withdraw()
        time.sleep(1)  # Wait for 1 second

        # Capture the screenshot of the entire screen
        screenshot = ImageGrab.grab()

        # Calculate the coordinates for the custom capture
        left = min(self.x1_coordinate, x2_coordinate)
        top = min(self.y1_coordinate, y2_coordinate)
        right = max(self.x1_coordinate, x2_coordinate)
        bottom = max(self.y1_coordinate, y2_coordinate)

        # Crop the screenshot to the custom capture area
        screenshot = screenshot.crop((left, top, right, bottom))

        # Save the screenshot to a file
        screenshot_path = f"{self.path_var.get()}Custom_Screenshot.png"
        screenshot.save(screenshot_path)

        # Add the screenshot to the Word document
        self.document.add_picture(screenshot_path)

        # Prompt for input comments if the checkbox is checked
        if self.comment_var.get():
            comment = simpledialog.askstring("Input", "Enter your comment:")
            if comment:
                self.document.add_paragraph(comment)

        # Save the Word document
        screenshot_docx_path = f"{self.path_var.get()}screenshot.docx"
        self.document.save(screenshot_docx_path)

        print(f"Custom screenshot captured and added to {screenshot_docx_path}")

        # Update the path label
        self.path_label.config(text=f"Path: {screenshot_docx_path}")

        # Show the window again
        self.root.deiconify()

    def update_coordinates(self):
        # Get the current x and y coordinates of the mouse cursor relative to the entire screen
        self.x_screen, self.y_screen = pyautogui.position()

        # Update labels with current mouse coordinates
        self.x_label.config(text=f"X Coordinate: {self.x_screen}")
        self.y_label.config(text=f"Y Coordinate: {self.y_screen}")

        # Capture the coordinates based on the current phase
        if self.capture_phase == 1:
            # Capture the coordinates for the first set
            self.x_entry.delete(0, tk.END)
            self.x_entry.insert(0, str(self.x_screen))
            self.y_entry.delete(0, tk.END)
            self.y_entry.insert(0, str(self.y_screen))

        # Schedule the next update after a delay (in milliseconds)
        self.root.after(100, self.update_coordinates)


if __name__ == "__main__":
    root = tk.Tk()
    capture_window = CaptureWindow(root)
    root.mainloop()
